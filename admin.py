from flask import Blueprint, render_template, jsonify, request
from db import db, User, PhishingURL
from flask import Blueprint, render_template, jsonify, request
from db import db, User, PhishingURL ,AuditLog, Url

from functools import wraps
import re
from flask import Flask, request, render_template, flash, redirect, url_for, session
from bs4 import BeautifulSoup
import requests
from urllib.parse import urljoin
from controller import Controller  
from flask_login import current_user



admin_bp = Blueprint('admin', __name__, url_prefix='/admin')



class LoggerService:
    @staticmethod
    def log_action(action, user_id, details=None):
        try:
            audit_log = AuditLog(
                action=action,
                user_id=user_id,
                details=details
            )
            db.session.add(audit_log)
            db.session.commit()
            return True
        except Exception as e:
            db.session.rollback()
            print(f"Error logging action: {str(e)}")
            return False 




from flask import render_template, request, redirect, url_for, session, flash
from datetime import datetime


# @admin_bp.route('/dashboard', methods=['GET', 'POST'])
# def dashboard():
#     if 'user_id' not in session:
#         return redirect(url_for('login'))

#     user = User.query.get(session['user_id'])
#     if user.role != 'admin':
#         flash('You do not have access to this page.', 'danger')
#         return redirect(url_for('index'))

#     # Get filter parameters from the request
#     action_filter = request.args.get('action')
#     user_filter = request.args.get('user')
#     date_filter = request.args.get('date')

#     # Initialize phishing URLs query
#     phishing_urls_query = PhishingURL.query

#     # Apply filters to phishing URLs query if provided
#     if action_filter:
#         phishing_urls_query = phishing_urls_query.filter(PhishingURL.status == action_filter)
#     if user_filter:
#         phishing_urls_query = phishing_urls_query.filter(PhishingURL.user.has(username=user_filter))
    
#     phishing_urls = phishing_urls_query.all()  # Fetch filtered phishing URLs

#     # Initialize audit logs query
#     audit_logs_query = AuditLog.query

#     # Apply filters to audit logs query if provided
#     if action_filter:
#         audit_logs_query = audit_logs_query.filter(AuditLog.action == action_filter)
#     if user_filter:
#         audit_logs_query = audit_logs_query.filter(AuditLog.user.has(username=user_filter))
#     if date_filter:
#         try:
#             date_object = datetime.strptime(date_filter, '%Y-%m-%d')
#             audit_logs_query = audit_logs_query.filter(AuditLog.timestamp.date() == date_object.date())
#         except ValueError:
#             flash('Invalid date format. Use YYYY-MM-DD.', 'danger')
#             return redirect(url_for('admin.dashboard'))

#     audit_logs = audit_logs_query.all()  # Fetch filtered audit logs

#     # Check if there are any results for phishing URLs and audit logs
#     phishing_urls_empty = not phishing_urls
#     audit_logs_empty = not audit_logs

#     # Render the template with the data and flags for no results
#     return render_template(
#         'admin/admin_dashboard.html', 
#         phishing_urls=phishing_urls, 
#         audit_logs=audit_logs,
#         phishing_urls_empty=phishing_urls_empty,
#         audit_logs_empty=audit_logs_empty
#     )


from sqlalchemy import func

@admin_bp.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if user.role != 'admin':
        flash('You do not have access to this page.', 'danger')
        return redirect(url_for('index'))

    # Get filter parameters from the request
    action_filter = request.args.get('action')
    user_filter = request.args.get('user')
    date_filter = request.args.get('date')

    # Initialize phishing URLs query
    phishing_urls_query = PhishingURL.query

    # Apply filters to phishing URLs query if provided
    if action_filter:
        phishing_urls_query = phishing_urls_query.filter(PhishingURL.status == action_filter)
    if user_filter:
        phishing_urls_query = phishing_urls_query.filter(PhishingURL.user.has(username=user_filter))
    
    phishing_urls = phishing_urls_query.all()  # Fetch filtered phishing URLs

    # Initialize audit logs query
    audit_logs_query = AuditLog.query

    # Apply filters to audit logs query if provided
    if action_filter:
        audit_logs_query = audit_logs_query.filter(AuditLog.action == action_filter)
    if user_filter:
        audit_logs_query = audit_logs_query.filter(AuditLog.user.has(username=user_filter))
    if date_filter:
        try:
            date_object = datetime.strptime(date_filter, '%Y-%m-%d')
            audit_logs_query = audit_logs_query.filter(AuditLog.timestamp.date() == date_object.date())
        except ValueError:
            flash('Invalid date format. Use YYYY-MM-DD.', 'danger')
            return redirect(url_for('admin.dashboard'))

    audit_logs = audit_logs_query.all()  # Fetch filtered audit logs

    # Calculate statistics
    custom_urls_count = db.session.query(func.count(Url.id)).filter(Url.custom_url.isnot(None)).scalar()
    short_urls_count = db.session.query(func.count(Url.id)).filter(Url.short_url.isnot(None)).scalar()
    total_clicks = db.session.query(func.sum(Url.clicks)).scalar()

    # Check if there are any results for phishing URLs and audit logs
    phishing_urls_empty = not phishing_urls
    audit_logs_empty = not audit_logs

    # Render the template with the data and flags for no results
    return render_template(
        'admin/admin_dashboard.html', 
        phishing_urls=phishing_urls, 
        audit_logs=audit_logs,
        phishing_urls_empty=phishing_urls_empty,
        audit_logs_empty=audit_logs_empty,
        custom_urls_count=custom_urls_count,
        short_urls_count=short_urls_count,
        total_clicks=total_clicks
    )



@admin_bp.route('/manage')
def manage():
    urls = PhishingURL.query.all()
    return render_template('admin/admin_manage.html', urls=urls)

@admin_bp.route('/users')
def users():
    user_list = User.query.all()
    return render_template('admin/admin_users.html', users=user_list)


import io
from docx import Document
from flask import render_template, send_file, flash, redirect, url_for
from datetime import datetime 
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile
import os

import io
import os
import tempfile
from flask import send_file
from docx import Document
from datetime import datetime
from werkzeug.utils import secure_filename


from flask import redirect, url_for, flash
from flask_login import logout_user

@admin_bp.route('/logout')
def logout():
    # Logout the user (clears the session)
    logout_user()

    # Clear any session data (if needed)
    session.clear()  # Uncomment this line if you are using Flask sessions and want to clear everything

    # Flash a message to inform the user 
    flash('You have been logged out.', 'info')

    # Redirect to the homepage or login page after logout
    return redirect(url_for('index'))  # Assuming 'index' is the name of the home route




@admin_bp.route('/reports')
def reports():
    # Get statistics
    custom_urls_count = db.session.query(func.count(Url.id)).filter(Url.custom_url.isnot(None)).scalar()
    short_urls_count = db.session.query(func.count(Url.id)).filter(Url.short_url.isnot(None)).scalar()
    total_clicks = db.session.query(func.sum(Url.clicks)).scalar()

    # Initialize phishing URLs and audit logs queries
    phishing_urls_query = PhishingURL.query
    audit_logs_query = AuditLog.query

    # Create the Word document
    doc = Document()

    # Add Title and Formatting (Professional, Bold Title)
    title = doc.add_heading('Phishing URL Report', 0)
    title.bold = True
    title.alignment = 1  # Center the title
    
    # Add Report Date (Smaller and lighter)
    doc.add_paragraph(f'Report Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', style='BodyText')
    doc.add_paragraph('')  # Add an empty line

    # Add Statistics Section (Using Table for Neatness)
    doc.add_heading('Statistics', level=1).bold = True
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = 'Table Grid'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Custom URLs Count'
    hdr_cells[1].text = str(custom_urls_count)
    stats_table.add_row().cells[0].text = 'Short URLs Count'
    stats_table.rows[1].cells[1].text = str(short_urls_count)
    stats_table.add_row().cells[0].text = 'Total Clicks'
    stats_table.rows[2].cells[1].text = str(total_clicks)

    doc.add_paragraph('')  # Add an empty line

    # Add Phishing URLs Data (Stylized for Better Readability)
    doc.add_heading('Phishing URLs', level=1).bold = True
    for url in phishing_urls_query.all():
        doc.add_paragraph(f'Status: {url.status}', style='BodyText').bold = True
        doc.add_paragraph(f'User: {url.user.username}', style='BodyText')
        doc.add_paragraph(f'Timestamp: {url.created_at}', style='BodyText')
        doc.add_paragraph('---', style='BodyText')

    # Add Audit Logs Data (Stylized with Bullet Points)
    doc.add_heading('Audit Logs', level=1).bold = True
    for log in audit_logs_query.all():
        doc.add_paragraph(f'Action: {log.action}', style='List Bullet').bold = True
        doc.add_paragraph(f'User: {log.user.username}', style='List Bullet')
        doc.add_paragraph(f'Timestamp: {log.timestamp}', style='List Bullet')
        doc.add_paragraph('---', style='BodyText')

    # Save the document to a BytesIO object
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    # Use tempfile to create a temporary .docx file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_doc_file:
        temp_doc_path = tmp_doc_file.name
        tmp_doc_file.write(doc_stream.read())

    # Ensure the file is saved
    if not os.path.exists(temp_doc_path):
        raise Exception("Failed to save the .docx file.")

    # Convert .docx to PDF using python-docx and pdfkit or other library if needed
    output_pdf_path = temp_doc_path.replace('.docx', '.pdf')
    try:
        # Implement the conversion logic here using your preferred library (e.g., pdfkit)
        convert_docx_to_pdf(temp_doc_path, output_pdf_path)
    except Exception as e:
        raise Exception(f"Failed to convert .docx to PDF: {e}")

    # Return the generated PDF as a downloadable file
    return send_file(
        output_pdf_path,
        as_attachment=True,
        download_name="phishing_url_report.pdf",
        mimetype='application/pdf'
    )

def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Converts a .docx file to a .pdf file using ReportLab.
    """
    doc = Document(docx_path)
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, height - 50, "Phishing URL Report")

    y_position = height - 100

    # Add contents from the DOCX file to PDF
    for para in doc.paragraphs:
        c.setFont("Helvetica", 10)
        text = para.text
        c.drawString(100, y_position, text)

        # Move to the next line
        y_position -= 15
        if y_position < 50:
            c.showPage()  # Add a new page if the text exceeds the page height
            c.setFont("Helvetica", 10)
            y_position = height - 50

    # Save the PDF file
    c.save()
 



def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if the user is logged in
        if 'user_id' not in session:
            return redirect(url_for('login'))  # Redirect to login if not authenticated

        # Fetch the user from the database
        user = User.query.get(session['user_id'])
        if not user:
            flash('User not found.', 'danger')
            return redirect(url_for('login'))  # Redirect to login if user doesn't exist

        # Check if the user is an admin
        if user.role != 'admin':
            flash('You do not have access to this page.', 'danger')
            return redirect(url_for('index'))  # Redirect to the home page or a "no access" page

        # Proceed with the original function if checks pass
        return f(*args, **kwargs)

    return decorated_function




# @admin_bp.route('/flag_url/<int:url_id>', methods=['POST'])
# def flag_url(url_id):
#     if 'user_id' not in session:
#         return redirect(url_for('login'))

#     user = User.query.get(session['user_id'])
#     if user.role != 'admin':
#         flash('You do not have access to this page.', 'danger')
#         return redirect(url_for('index'))

#     phishing_url = PhishingURL.query.get_or_404(url_id)
#     phishing_url.flagged = True
#     db.session.commit()

#     # Log the action in audit logs
#     LoggerService.log_action(
#             action=f"Flagged URL: {phishing_url.url}",
#             user_id=current_user.id,
#             details=f"URL ID: {url_id}, Status: Flagged"
#         )

#     flash(f'URL {phishing_url.url} flagged as phishing.', 'success')
#     return redirect(url_for('dashboard'))


@admin_bp.route('/f', methods=['GET', 'POST'])
def f():
    if request.method == 'POST':
        url = request.form['url']
        category = request.form['category']
        status = 'phishing'

        # Check if the URL already exists
        existing_url = PhishingURL.query.filter_by(url=url).first()
        if existing_url:
            flash(f'The URL {url} already exists.', 'danger')
            return redirect(url_for('admin.f'))

        # Insert new URL into the database
        new_url = PhishingURL(url=url, category=category, status=status, flagged=True, user_id=current_user.id)
        db.session.add(new_url)
        db.session.commit()

        flash(f'New URL {url} added successfully.', 'success')
        return redirect(url_for('admin.f'))

    return render_template('admin/insert_flag.html')


@admin_bp.route('/change_flag', methods=['GET', 'POST'])
def change_flag():
    if request.method == 'POST':
        url_id = request.form['url_id']
        url = PhishingURL.query.get_or_404(url_id)
        
        # Toggle flag status
        url.flagged = not url.flagged
        db.session.commit()

        flash(f'URL {url.url} {"Flagged" if url.flagged else "Unflagged"} successfully.', 'success')
        return redirect(url_for('admin.change_flag'))

    return render_template('admin/change_flag.html', urls=PhishingURL.query.all())



@admin_bp.route('/flag_url/<int:url_id>', methods=['POST'])
def flag_url(url_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if user.role != 'admin':
        flash('You do not have access to this page.', 'danger')
        return redirect(url_for('index'))

    phishing_url = PhishingURL.query.get_or_404(url_id)
    phishing_url.flagged = not phishing_url.flagged  # Toggle flag status
    db.session.commit()

    # Log the action in audit logs for flagging/unflagging
    LoggerService.log_action(
        action=f"{'Flagged' if phishing_url.flagged else 'Unflagged'} URL: {phishing_url.url}",
        user_id=user.id,
        details=f"URL ID: {url_id}, Status: {'Flagged' if phishing_url.flagged else 'Unflagged'}"
    )

    flash(f'URL {phishing_url.url} {"flagged" if phishing_url.flagged else "unflagged"}.', 'success')
    return redirect(url_for('admin.change_flag'))


def log_action(action, user_id):
    log = AuditLog(action=action, user_id=user_id)
    db.session.add(log)
    db.session.commit()




@admin_bp.route('/audit_logs')
@admin_required
def audit_logs():
    page = request.args.get('page', 1, type=int)
    logs = AuditLog.query.order_by(AuditLog.timestamp.desc())\
        .paginate(page=page, per_page=20)
    return render_template('admin/audit_logs.html', logs=logs)



@admin_bp.route('/manage_users')
def manage_users():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    if user.role != 'admin':
        flash('You do not have access to this page.', 'danger')
        return redirect(url_for('index'))

    users = User.query.all()  # Get all users
    return render_template('admin/manage_users.html', users=users)

@admin_bp.route('/change_role/<int:user_id>', methods=['POST'])
def change_role(user_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    admin = User.query.get(session['user_id'])
    if admin.role != 'admin':
        flash('You do not have access to this page.', 'danger')
        return redirect(url_for('index'))

    user = User.query.get_or_404(user_id)
    new_role = request.form.get('role')
    user.role = new_role
    db.session.commit()

    # Log the role change in audit logs
    log_action(f"Changed role of {user.username} to {new_role}", admin.id)

    flash(f'Role of {user.username} changed to {new_role}.', 'success')
    return redirect(url_for('admin.manage_users'))


@admin_bp.route('/toggle_active/<int:user_id>', methods=['POST'])
def toggle_active(user_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    admin = User.query.get(session['user_id'])
    if admin.role != 'admin':
        flash('You do not have access to this page.', 'danger')
        return redirect(url_for('index'))

    user = User.query.get_or_404(user_id)
    user.active = not user.active  # Toggle active status
    db.session.commit()

    # Log the action
    LoggerService.log_action(
        action=f"{'Activated' if user.active else 'Deactivated'} user: {user.username}",
        user_id=admin.id,
        details=f"User ID: {user_id}, Role: {user.role}"
    )

    flash(f"User {user.username} has been {'activated' if user.active else 'deactivated'}.", 'success')
    return redirect(url_for('admin.manage_users'))



@admin_bp.route('/edit_user/<int:user_id>', methods=['GET', 'POST'])
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)
    current_user_id = session.get('user_id')
    
    if request.method == 'POST':
        try:
            # Get form data
            username = request.form['username']
            email = request.form['email']
            role = request.form['role']
            
            # Validate Username: Must be alphanumeric and between 3–50 characters
             # Validate Username: Must be alphabetic and between 3–50 characters
            if not username.isalpha():
                flash('Username must contain only alphabetic characters (A-Z, a-z).', 'danger')
                return redirect(url_for('admin.edit_user', user_id=user_id))

            if len(username) < 3 or len(username) > 50:
                flash('Username must be between 3 and 50 characters long.', 'danger')
                return redirect(url_for('admin.edit_user', user_id=user_id))

            # Validate Email: Check if it's in the correct format
            if not re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', email):
                flash('Please provide a valid email address.', 'danger')
                return redirect(url_for('admin.edit_user', user_id=user_id))
            
            # Check if the email already exists in the database
            existing_user = User.query.filter_by(email=email).first()
            if existing_user and existing_user.id != user_id:
                flash('Email address is already in use by another user.', 'danger')
                return redirect(url_for('admin.edit_user', user_id=user_id))

            # Validate Role: Must be either 'admin' or 'user'
            if role not in ['admin', 'user']:
                flash('Invalid role selected.', 'danger')
                return redirect(url_for('admin.edit_user', user_id=user_id))

            # Update user details
            user.username = username
            user.email = email
            user.role = role
            
            # Commit the changes to the database
            db.session.commit()

            # Log the action
            LoggerService.log_action(
                action=f"Updated user: {user.username}",
                user_id=current_user_id,
                details=f"User ID: {user_id}, Role: {user.role}"
            )

            flash('User updated successfully.', 'success')
            return redirect(url_for('admin.manage_users'))
        
        except Exception as e:
            db.session.rollback()  # Rollback the transaction in case of an error
            flash(f'Error updating user: {str(e)}', 'danger')
            return redirect(url_for('admin.edit_user', user_id=user_id))

    return render_template('admin/edit_users.html', user=user)